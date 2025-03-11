import argparse
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def create_document(path: str, output: str = "output.docx") -> None:
    doc = Document()

    with open(path, "r", encoding="utf-8") as f: content = f.read().split("\n")

    headers = []
    contents = []
    tables = []
    
    content_buffer = []

    for line in content:
        line = line.strip()

        if line.lower().startswith("header:"):
            headers.append(line[len("header:"):].strip())

        elif line.lower().startswith("content:"):
            content_buffer.append(line[len("content:"):].strip())

            if len(content_buffer) >= len(headers):
                contents.append(content_buffer)
                content_buffer = []

        elif line.lower() == "endtable":
            if headers and contents:
                tables.append((headers[:], contents[:]))
                
            headers.clear()
            contents.clear()
            content_buffer.clear()

    if len(content_buffer) >= len(headers): 
        contents.append(content_buffer)

    for headers, contents in tables:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        header_row = table.rows[0].cells
        for i, header in enumerate(headers):
            header_row[i]._element.get_or_add_tcPr().append(
                parse_xml(r'<w:shd {} w:fill="4287f5"/>'.format(nsdecls('w')))
            )

            paragraph = header_row[i].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            run = paragraph.add_run(header)
            run.font.name = "Arial"
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(255, 255, 255)

        for row_data in contents:
            row = table.add_row().cells

            for i, cell_data in enumerate(row_data):
                paragraph = row[i].paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                run = paragraph.add_run(cell_data)
                run.font.name = "Arial"
                run.font.size = Pt(11)

        doc.add_paragraph("")

    doc.save(output)
    print(f"Document saved as {output}!")

def main() -> None:
    parser = argparse.ArgumentParser(description="Process docx tables")
    parser.add_argument("--path", type=str, required=True, help="File path")
    parser.add_argument("--output", type=str, required=False, default="output.docx", help="File output path")

    args = parser.parse_args()
    create_document(args.path, args.output)

if __name__ == "__main__":
    main()